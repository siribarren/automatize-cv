# src/generators.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def agregar_titulo(doc, texto):
    p = doc.add_heading(texto, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def agregar_parrafo(doc, texto, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = bold
    run.font.size = Pt(11)

def guardar_docx_output(respuesta, output_folder="data/output", nombre="candidato"):
    os.makedirs(output_folder, exist_ok=True)

    # Separar bloques por encabezados
    secciones = separar_por_bloques(respuesta)

    # === Generar CV ===
    doc_cv = Document()
    agregar_titulo(doc_cv, "Resumen Profesional")
    agregar_parrafo(doc_cv, secciones.get("Resumen profesional", "Sin resumen"))

    agregar_titulo(doc_cv, "Experiencia Laboral")
    agregar_parrafo(doc_cv, secciones.get("Experiencia laboral", "No especificada"))

    if "Tabla de habilidades" in secciones:
        agregar_titulo(doc_cv, "Tabla de Habilidades")
        agregar_tabla_habilidades(doc_cv, secciones["Tabla de habilidades"])

    cv_path = os.path.join(output_folder, f"cv_{nombre}.docx")
    doc_cv.save(cv_path)

    # === Generar Carta de Presentación ===
    carta_texto = secciones.get("Carta de presentación", None)
    if carta_texto:
        doc_carta = Document()
        agregar_parrafo(doc_carta, carta_texto)
        carta_path = os.path.join(output_folder, f"carta_presentacion_{nombre}.docx")
        doc_carta.save(carta_path)

    return output_folder  # Puedes luego empaquetar los archivos si deseas
